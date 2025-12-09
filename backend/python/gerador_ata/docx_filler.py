"""
DOCX Filler - Preenche template DOCX com dados estruturados
SDC-Ata-Generator
"""

import os
from datetime import datetime
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

try:
    from . import config_ata as config
except ImportError:
    import config_ata as config


class DOCXFillerError(Exception):
    """Exceção customizada para erros no preenchimento DOCX"""
    pass


class DOCXFiller:
    """Classe para preencher templates DOCX com dados estruturados"""

    def __init__(self, template_path: str = None):
        """
        Inicializa o preenchedor DOCX

        Args:
            template_path: Caminho para o arquivo template (opcional, usa config se não fornecido)
        """
        self.template_path = template_path or config.TEMPLATE_PATH

        if not os.path.exists(self.template_path):
            raise DOCXFillerError(f"Template não encontrado: {self.template_path}")

        self.doc = None

    def fill_template(self, data: Dict[str, Any], output_filename: str = None) -> str:
        """
        Preenche o template com os dados e salva o arquivo

        Args:
            data: Dicionário com dados estruturados
            output_filename: Nome do arquivo de saída (opcional)

        Returns:
            Caminho completo do arquivo gerado

        Raises:
            DOCXFillerError: Se houver erro no preenchimento
        """
        try:
            # Carregar template
            self.doc = Document(self.template_path)

            # Validar dados
            self._validate_data(data)

            # Preencher documento (ordem correta: Header → Participantes → Objetivo → Pontos → Próximos Passos)
            self._fill_header_table(data)
            self._fill_participants_table(data)
            self._fill_objective_paragraph(data)
            self._fill_topics_table(data)
            self._fill_next_steps_table(data)

            # Gerar nome do arquivo de saída
            if not output_filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"Ata_Reuniao_{timestamp}.docx"

            # Garantir extensão .docx
            if not output_filename.endswith('.docx'):
                output_filename += '.docx'

            # Caminho completo de saída
            output_path = os.path.join(config.OUTPUT_DIR, output_filename)

            # Salvar documento
            self.doc.save(output_path)

            return output_path

        except Exception as e:
            raise DOCXFillerError(f"Erro ao preencher template: {str(e)}")

    def _validate_data(self, data: Dict[str, Any]):
        """Valida se os dados contêm todos os campos necessários"""
        required_fields = ['local', 'data_horario', 'convocado_por', 'objetivo']
        required_arrays = ['participantes', 'pontos', 'proximos_passos']

        missing = []

        for field in required_fields:
            if field not in data or not data[field]:
                missing.append(field)

        for array in required_arrays:
            if array not in data or not isinstance(data[array], list) or len(data[array]) == 0:
                missing.append(array)

        if missing:
            raise DOCXFillerError(f"Campos obrigatórios faltando: {', '.join(missing)}")

    def _fill_header_table(self, data: Dict[str, Any]):
        """Preenche a tabela de cabeçalho (Tabela 1)"""
        if len(self.doc.tables) < 1:
            raise DOCXFillerError("Template não contém tabela de cabeçalho")

        table = self.doc.tables[0]

        # Verificar se a tabela tem o formato esperado (2 linhas: header + dados)
        if len(table.rows) < 2:
            raise DOCXFillerError("Tabela de cabeçalho com formato inválido")

        # Preencher segunda linha (dados)
        data_row = table.rows[1]
        if len(data_row.cells) >= 3:
            data_row.cells[0].text = data['local']
            data_row.cells[1].text = data['data_horario']
            data_row.cells[2].text = data['convocado_por']

    def _reorganize_document_elements(self):
        """
        Reorganiza os elementos do documento (MÉTODO LEGADO - NÃO USADO).
        A ordem correta agora é garantida pela sequência de chamadas em fill_template().
        Ordem esperada: Cabeçalho → Participantes → Objetivo → Pontos → Próximos Passos
        """
        body = self.doc.element.body

        # Encontrar os elementos por tipo
        header_table = None
        participants_table = None
        objective_paras = []

        table_count = 0
        for element in body:
            if element.tag.endswith('tbl'):
                table_count += 1
                if table_count == 1:
                    header_table = element
                elif table_count == 2:
                    participants_table = element
            elif element.tag.endswith('p'):
                text = ''.join([t.text for t in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')])
                if 'Objetivo' in text:
                    objective_paras.append(element)

        # Se encontrou Participantes e Objetivo, verificar ordem
        if participants_table is not None and objective_paras:
            # Obter índices na ordem do body
            elements = list(body)
            try:
                participants_idx = elements.index(participants_table)
                objective_idx = elements.index(objective_paras[0])

                # Se Objetivo vem antes de Participantes, reordenar
                if objective_idx < participants_idx:
                    # Remover a tabela Participantes de sua posição
                    body.remove(participants_table)

                    # Encontrar a nova posição (deve ser após Header, antes de Objetivo)
                    # Recalcular índices após remoção
                    elements = list(body)
                    if header_table is not None:
                        header_idx = elements.index(header_table)
                        # Inserir após o Header
                        header_table.addnext(participants_table)
            except ValueError:
                # Se não encontrar, apenas pular
                pass

    def _fill_objective_paragraph(self, data: Dict[str, Any]):
        """
        Preenche o parágrafo com o objetivo da reunião

        Busca pelo placeholder específico '[Descrição breve do objetivo da reunião.]'
        e o substitui mantendo o título "Objetivo" intacto.
        """
        placeholder_objetivo = '[Descrição breve do objetivo da reunião.]'

        for paragraph in self.doc.paragraphs:
            if placeholder_objetivo in paragraph.text:
                # Substituir apenas o placeholder, mantendo o resto do parágrafo
                paragraph.text = paragraph.text.replace(placeholder_objetivo, data['objetivo'])
                break

    def _preencher_tabela_preservando_formato(self, tabela, dados_lista, campos):
        """
        Preenche tabela copiando formatação da linha modelo
        Usa addnext() para manter a ordem dos elementos no documento

        Args:
            tabela: Objeto Table
            dados_lista: Lista de dicionários
            campos: Lista de campos na ordem das colunas
        """
        if not dados_lista or len(tabela.rows) < 2:
            return

        # Primeira entrada: preencher linha modelo existente (índice 1)
        linha = tabela.rows[1]
        for idx, campo in enumerate(campos):
            if idx < len(linha.cells):
                linha.cells[idx].text = str(dados_lista[0].get(campo, ''))

        # Demais entradas: copiar linha modelo e inserir após a última
        ultima_linha_idx = 1
        for i in range(1, len(dados_lista)):
            item = dados_lista[i]

            # Copiar XML da linha modelo (sempre linha 1)
            tr_modelo = tabela.rows[1]._tr
            new_tr = deepcopy(tr_modelo)

            # CORREÇÃO: Inserir APÓS a última linha preenchida (mantém ordem)
            # addnext() insere imediatamente após, sem mover a tabela
            tabela.rows[ultima_linha_idx]._tr.addnext(new_tr)

            # Atualizar índice e preencher nova linha
            ultima_linha_idx += 1
            nova_linha = tabela.rows[ultima_linha_idx]

            for idx, campo in enumerate(campos):
                if idx < len(nova_linha.cells):
                    nova_linha.cells[idx].text = str(item.get(campo, ''))

    def _fill_table_with_formatting(self, table, data_list, fields):
        """
        Preenche tabela copiando a formatação da linha modelo para cada nova linha

        Args:
            table: Objeto Table do python-docx
            data_list: Lista de dicionários com os dados
            fields: Lista de nomes dos campos na ordem das colunas
        """
        if not data_list or len(table.rows) < 2:
            return

        # A linha 1 (índice 1) é o modelo - salvamos antes de remover
        model_row = table.rows[1]

        # Para cada item de dados, copiar a linha modelo e preencher
        for i, item in enumerate(data_list):
            if i == 0:
                # Primeira entrada: usar a linha modelo existente (primeira linha de dados)
                # Após remover linhas, a linha 1 passa a ser índice 1 no caso de ser a primeira a preencher
                row = table.rows[1]
            else:
                # Demais entradas: copiar linha modelo com formatação
                row = self._copy_row_formatting(table, 1)

            # Preencher células
            for idx, field in enumerate(fields):
                if idx < len(row.cells):
                    valor = str(item.get(field, ''))
                    row.cells[idx].text = valor

    def _fill_participants_table(self, data: Dict[str, Any]):
        """Preenche a tabela de participantes (Tabela 2)"""
        if len(self.doc.tables) < 2:
            raise DOCXFillerError("Template não contém tabela de participantes")

        table = self.doc.tables[1]

        # Verificar se tem header (linha 0) e modelo (linha 1)
        if len(table.rows) < 2:
            raise DOCXFillerError("Tabela de participantes inválida - precisa de header e modelo")

        # Remover linhas de dados existentes (manter apenas header e modelo)
        while len(table.rows) > 2:
            table._element.remove(table.rows[-1]._element)

        # Preencher com formatação preservada
        self._preencher_tabela_preservando_formato(
            table,
            data['participantes'],
            ['num', 'nome']
        )

    def _fill_topics_table(self, data: Dict[str, Any]):
        """Preenche a tabela de pontos discutidos (Tabela 3)"""
        if len(self.doc.tables) < 3:
            raise DOCXFillerError("Template não contém tabela de pontos discutidos")

        table = self.doc.tables[2]

        # Remover linhas de dados existentes (manter apenas header e modelo)
        while len(table.rows) > 2:
            table._element.remove(table.rows[-1]._element)

        # Preencher com formatação preservada
        self._preencher_tabela_preservando_formato(
            table,
            data['pontos'],
            ['item', 'topico']
        )

    def _fill_next_steps_table(self, data: Dict[str, Any]):
        """Preenche a tabela de próximos passos (Tabela 4)"""
        if len(self.doc.tables) < 4:
            raise DOCXFillerError("Template não contém tabela de próximos passos")

        table = self.doc.tables[3]

        # Remover linhas de dados existentes (manter apenas header e modelo)
        while len(table.rows) > 2:
            table._element.remove(table.rows[-1]._element)

        # Preencher com formatação preservada
        self._preencher_tabela_preservando_formato(
            table,
            data['proximos_passos'],
            ['item', 'acao', 'responsavel', 'data']
        )


def fill_docx_from_data(data: Dict[str, Any], output_filename: str = None, template_path: str = None) -> str:
    """
    Função utilitária para preencher DOCX a partir de dados

    Args:
        data: Dicionário com dados estruturados
        output_filename: Nome do arquivo de saída (opcional)
        template_path: Caminho do template (opcional)

    Returns:
        Caminho completo do arquivo gerado
    """
    filler = DOCXFiller(template_path)
    return filler.fill_template(data, output_filename)


# Exemplo de uso
if __name__ == '__main__':
    # Dados de exemplo
    exemplo_dados = {
        "local": "Sala de Reuniões SDC",
        "data_horario": "13/11/2025 - 14:00 às 15:30",
        "convocado_por": "Jonathan Silva",
        "objetivo": "Demonstração de processo para geração de atas via transcrição automática",
        "participantes": [
            {"num": "1", "nome": "Jonathan Silva"},
            {"num": "2", "nome": "Everton Santos"},
            {"num": "3", "nome": "Maria Costa"}
        ],
        "pontos": [
            {"item": "1", "topico": "Gravação utilizando OBS configurado para salvar vídeos em formato MP4"},
            {"item": "2", "topico": "Compressão de arquivo reduzindo de 956MB para 69MB usando Clipchamp"},
            {"item": "3", "topico": "Transcrição automática via Word Online com limite de 300MB/mês"}
        ],
        "proximos_passos": [
            {
                "item": "1",
                "acao": "Padronizar procedimento para todas as gravações internas",
                "responsavel": "Jonathan Silva",
                "data": "20/11/2025"
            },
            {
                "item": "2",
                "acao": "Compartilhar tutorial com colaboradores",
                "responsavel": "Everton Santos",
                "data": "25/11/2025"
            }
        ]
    }

    try:
        print("Preenchendo template DOCX...")
        output_path = fill_docx_from_data(exemplo_dados, "Ata_Teste.docx")
        print(f"Arquivo gerado com sucesso: {output_path}")
    except DOCXFillerError as e:
        print(f"Erro ao gerar arquivo: {e}")
